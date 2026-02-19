import os
import logging
import io
from flask import Flask, render_template, request, jsonify, session, send_file
from flask_session import Session
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib import colors
from reportlab.lib.units import inch
from reportlab.platypus import Preformatted
from docx import Document
import google.generativeai as genai

# ---------------------------------------------------
# FLASK SETUP
# ---------------------------------------------------

app = Flask(__name__)
app.secret_key = os.environ.get("SECRET_KEY", "supersecretkey")

app.config["SESSION_PERMANENT"] = False
app.config["SESSION_TYPE"] = "filesystem"
Session(app)

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# ---------------------------------------------------
# GEMINI SETUP
# ---------------------------------------------------

genai.configure(api_key=os.environ.get("GEMINI_API_KEY"))
model = genai.GenerativeModel("models/gemini-1.5-flash")

def query_ai(prompt, system_prompt=""):
    try:
        full_prompt = f"{system_prompt}\n\n{prompt}"
        response = model.generate_content(full_prompt)
        return response.text
    except Exception as e:
        logger.error(f"Gemini error: {e}")
        return None

# ---------------------------------------------------
# ROUTES
# ---------------------------------------------------

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/generate_screenplay', methods=['POST'])
def generate_screenplay():
    data = request.json
    story = data.get('story')
    genre = data.get('genre')
    language = data.get('language')

    system_prompt = f"""
You are a professional Hollywood screenwriter.
Write a cinematic screenplay in {genre.upper()} style.
Language: {language}
Use:
- Scene headings in ALL CAPS
- Character names centered
- Strong pacing
- Emotional depth
- Visual storytelling
"""

    prompt = f"Story idea: {story}"

    response = query_ai(prompt, system_prompt)

    if not response:
        response = "Error generating screenplay."

    session['screenplay'] = response
    session['story'] = story
    session['genre'] = genre

    return jsonify({"content": response})

@app.route('/generate_characters', methods=['POST'])
def generate_characters():
    story = session.get('story', '')

    if not story:
        return jsonify({"error": "No story found"}), 400

    system_prompt = """
You are a character psychologist.
Create deep character profiles including:
- Age
- Background
- Motivation
- Internal conflict
- Fear
- Moral flaw
"""

    response = query_ai(f"Story: {story}", system_prompt)

    session['characters'] = response
    return jsonify({"content": response})

@app.route('/generate_director_mode', methods=['POST'])
def generate_director_mode():
    screenplay = session.get('screenplay', '')

    if not screenplay:
        return jsonify({"error": "No screenplay found"}), 400

    system_prompt = """
You are a film director.
Create a shot breakdown for each scene including:
- Camera angle
- Shot type
- Lighting style
- Color grading
- Emotional tone
"""

    response = query_ai(f"Screenplay: {screenplay}", system_prompt)

    session['director_mode'] = response
    return jsonify({"content": response})

@app.route('/generate_pitch_deck', methods=['POST'])
def generate_pitch_deck():
    story = session.get('story', '')

    if not story:
        return jsonify({"error": "No story found"}), 400

    system_prompt = """
You are a Hollywood producer.
Create a professional pitch including:
- Logline
- Tagline
- Genre
- Target audience
- Comparable films
- Why it will succeed
"""

    response = query_ai(f"Story: {story}", system_prompt)

    session['pitch_deck'] = response
    return jsonify({"content": response})

@app.route('/generate_sound_design', methods=['POST'])
def generate_sound_design():
    screenplay = session.get('screenplay', '')

    if not screenplay:
        return jsonify({"error": "No screenplay found"}), 400

    system_prompt = """
You are a professional sound designer.
Create a sound design plan including:
- Background music
- Ambient layers
- Foley effects
- Dialogue treatment
"""

    response = query_ai(f"Screenplay: {screenplay}", system_prompt)

    session['sound_design'] = response
    return jsonify({"content": response})

@app.route('/export/<format_type>')
def export_file(format_type):

    content = ""
    if 'screenplay' in session:
        content += f"\nSCREENPLAY\n\n{session['screenplay']}\n\n"
    if 'characters' in session:
        content += f"\nCHARACTERS\n\n{session['characters']}\n\n"
    if 'director_mode' in session:
        content += f"\nDIRECTOR NOTES\n\n{session['director_mode']}\n\n"
    if 'sound_design' in session:
        content += f"\nSOUND DESIGN\n\n{session['sound_design']}\n\n"
    if 'pitch_deck' in session:
        content += f"\nPITCH DECK\n\n{session['pitch_deck']}\n\n"

    if not content:
        return "No content to export", 400

    if format_type == "pdf":
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        elements = []
        style = ParagraphStyle(name='Normal', fontSize=10)
        elements.append(Preformatted(content, style))
        doc.build(elements)
        buffer.seek(0)

        return send_file(
            buffer,
            as_attachment=True,
            download_name="project_bible.pdf",
            mimetype="application/pdf"
        )

    elif format_type == "docx":
        doc = Document()
        doc.add_heading("Project Bible", 0)
        doc.add_paragraph(content)

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        return send_file(
            buffer,
            as_attachment=True,
            download_name="project_bible.docx",
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    return "Invalid format", 400

# ---------------------------------------------------

if __name__ == '__main__':
    app.run(debug=True)
